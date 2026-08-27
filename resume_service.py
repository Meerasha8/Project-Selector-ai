import json
import os
import tempfile
from typing import Any

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from fastapi import HTTPException
from groq import Groq
from pydantic import BaseModel

from embedding_utils import embed_text, cosine_similarity


class ResumeContent(BaseModel):
    summary: str
    skills: list[str]
    experience: list[dict[str, Any]]
    projects: list[dict[str, Any]]
    education: list[dict[str, Any]]
    certificates: list[dict[str, Any]]


class ResumeService:
    def __init__(self, api_key: str | None = None):
        self.client = Groq(api_key=api_key) if api_key else None
        if self.client is None:
            self.client = Groq(api_key=os.getenv("GROQ_API_KEY")) if os.getenv("GROQ_API_KEY") else None

    def _create_fallback_content(self, user_data: dict[str, Any]) -> ResumeContent:
        user_details = user_data.get("user_details", {})
        summary = user_details.get("profession_summary") or "Experienced software engineering professional."
        skills = [s.get("name") for s in user_data.get("skills", []) if s.get("name")]
        
        experience = []
        for item in user_data.get("internship", []):
            experience.append({
                "role": item.get("role") or "Intern",
                "company": item.get("company_name") or "Company",
                "duration": item.get("duration") or item.get("Duration") or "",
                "highlights": [item.get("description")] if item.get("description") else []
            })
            
        projects = []
        for item in user_data.get("projects", []):
            projects.append({
                "name": item.get("name") or "Project",
                "highlights": [f"Tech: {item.get('tech_stack')}", item.get("description")] if item.get("description") else []
            })

        education = user_data.get("education", [])
        certificates = user_data.get("certificates", [])

        return ResumeContent(
            summary=summary,
            skills=skills or ["Software Development"],
            experience=experience,
            projects=projects,
            education=education,
            certificates=certificates
        )

    def _filter_user_data_with_rag(self, job_description: str, user_data: dict[str, Any]) -> dict[str, Any]:
        job_embedding = embed_text(job_description)
        filtered_data = {"user_details": user_data.get("user_details", {})}

        def rank_and_filter(items, text_extractor, top_k):
            if not items:
                return []
            scored_items = []
            for item in items:
                text = text_extractor(item)
                emb = embed_text(text)
                score = cosine_similarity(job_embedding, emb)
                scored_items.append((score, item))
            scored_items.sort(key=lambda x: x[0], reverse=True)
            return [item for score, item in scored_items[:top_k]]

        filtered_data["skills"] = rank_and_filter(
            user_data.get("skills", []),
            lambda x: f"{x.get('name', '')} {x.get('description', '')}",
            top_k=15
        )

        filtered_data["internship"] = rank_and_filter(
            user_data.get("internship", []),
            lambda x: f"{x.get('role', '')} {x.get('description', '')}",
            top_k=5
        )

        filtered_data["projects"] = rank_and_filter(
            user_data.get("projects", []),
            lambda x: f"{x.get('name', '')} {x.get('description', '')} {x.get('tech_stack', '')}",
            top_k=4
        )

        filtered_data["education"] = rank_and_filter(
            user_data.get("education", []),
            lambda x: f"{x.get('course_name', '')} {x.get('college_name', '')}",
            top_k=3
        )

        filtered_data["certificates"] = rank_and_filter(
            user_data.get("certificates", []),
            lambda x: f"{x.get('certificate_name', '')}",
            top_k=5
        )

        return filtered_data

    def select_resume_content(self, job_description: str, user_data: dict[str, Any]) -> ResumeContent:
        if self.client is None:
            return self._create_fallback_content(user_data)

        filtered_user_data = self._filter_user_data_with_rag(job_description, user_data)
        raw_payload = json.dumps(filtered_user_data, ensure_ascii=False, indent=2)
        schema = {
            "summary": "string",
            "skills": ["string"],
            "experience": [
                {"role": "string", "company": "string", "duration": "string", "highlights": ["string"]}
            ],
            "projects": [{"name": "string", "highlights": ["string"]}],
            "education": [
                {
                    "course_name": "string",
                    "college_name": "string",
                    "location": "string",
                    "start_year": "number",
                    "end_year": "number",
                    "cgpa": "number",
                }
            ],
            "certificates": [{"certificate_name": "string", "certificate_issuer": "string"}],
        }
        prompt = (
            "You are selecting content for a one-page resume tailored to a job description. "
            "Use only the provided user data. Select only the most relevant projects, skills, experience, education, and certificates. "
            "Rewrite bullets concisely, action-verb-first, and quantified where possible. "
            "Return strict raw JSON matching this schema: " + json.dumps(schema) + ". "
            "Do not wrap in markdown or prose outside the JSON."
        )

        for attempt in range(3):
            try:
                response = self.client.chat.completions.create(
                    model=os.getenv("GROQ_CHAT_MODEL", "llama-3.3-70b-versatile"),
                    messages=[
                        {"role": "system", "content": "Return only raw valid JSON matching the requested schema. Do not use markdown codeblocks."},
                        {
                            "role": "user",
                            "content": f"Job description:\n{job_description}\n\nUser data:\n{raw_payload}\n\nInstructions:\n{prompt}",
                        },
                    ],
                    temperature=0.2,
                    max_tokens=1200,
                )
                content = (response.choices[0].message.content or "{}").strip()
                if content.startswith("```"):
                    lines = content.splitlines()
                    if lines[0].startswith("```"):
                        lines = lines[1:]
                    if lines and lines[-1].startswith("```"):
                        lines = lines[:-1]
                    content = "\n".join(lines).strip()
                
                parsed = json.loads(content)
                return ResumeContent(**parsed)
            except Exception as e:
                print(f"Groq resume parsing attempt {attempt} error: {e}")
                if attempt == 2:
                    break

        return self._create_fallback_content(user_data)

    def render_docx(self, content: ResumeContent, user_details: dict[str, Any] | None = None, template_style: str = "modern") -> bytes:
        template_style = (template_style or "modern").lower()
        if template_style == "classic":
            return self._render_classic_docx(content, user_details)
        elif template_style == "tech":
            return self._render_tech_docx(content, user_details)
        else:
            return self._render_modern_docx(content, user_details)

    def render_pdf(self, content: ResumeContent, user_details: dict[str, Any] | None = None, template_style: str = "minimalist") -> bytes:
        template_style = (template_style or "minimalist").lower()
        if "two" in template_style or "classic" in template_style:
            tmpl_name = "jinja_twocolumn.html"
        elif "bold" in template_style or "header" in template_style or "creative" in template_style:
            tmpl_name = "jinja_boldheader.html"
        else:
            tmpl_name = "jinja_minimalist.html"

        tmpl_path = os.path.join(os.path.dirname(__file__), "templates", tmpl_name)
        if not os.path.exists(tmpl_path):
            tmpl_path = os.path.join(os.path.dirname(__file__), tmpl_name)

        with open(tmpl_path, "r", encoding="utf-8") as f:
            template_str = f.read()

        try:
            from jinja2 import Template
            jinja_tmpl = Template(template_str)
            rendered_html = jinja_tmpl.render(
                user_details=user_details or {},
                summary=content.summary,
                skills=content.skills,
                experience=content.experience,
                projects=content.projects,
                education=content.education,
                certificates=content.certificates,
            )
        except Exception:
            rendered_html = template_str

        try:
            from weasyprint import HTML
            return HTML(string=rendered_html).write_pdf()
        except Exception as e:
            print(f"WeasyPrint PDF rendering fallback to HTML bytes: {e}")
            return rendered_html.encode("utf-8")

    # Template 1: Modern Clean ATS (Calibri / Arial, Sleek Dark Slate accents)
    def _render_modern_docx(self, content: ResumeContent, user_details: dict[str, Any] | None = None) -> bytes:
        document = Document()
        for section in document.sections:
            section.top_margin = Inches(0.4)
            section.bottom_margin = Inches(0.4)
            section.left_margin = Inches(0.45)
            section.right_margin = Inches(0.45)

        style = document.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(10.5)

        title = document.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run((user_details or {}).get("name") or "Resume")
        run.bold = True
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(15, 23, 42)

        contact = document.add_paragraph()
        contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_lines = []
        for key in ["email", "phone", "location", "github", "linkedin", "portfolio"]:
            val = (user_details or {}).get(key)
            if val:
                contact_lines.append(str(val))
        if contact_lines:
            c_run = contact.add_run(" | ".join(contact_lines))
            c_run.font.size = Pt(9.5)
            c_run.font.color.rgb = RGBColor(71, 85, 105)

        def add_heading(text):
            h = document.add_paragraph()
            h.paragraph_format.space_before = Pt(8)
            h.paragraph_format.space_after = Pt(2)
            r = h.add_run(text.upper())
            r.bold = True
            r.font.size = Pt(11)
            r.font.color.rgb = RGBColor(14, 116, 144)

        if content.summary:
            add_heading("Professional Summary")
            document.add_paragraph(content.summary)

        if content.skills:
            add_heading("Technical Skills")
            sp = document.add_paragraph()
            sp.add_run(", ".join(content.skills[:15]))

        if content.experience:
            add_heading("Work Experience")
            for item in content.experience[:4]:
                p = document.add_paragraph()
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(1)
                r_title = p.add_run(f"{item.get('role', '')} — {item.get('company', '')}")
                r_title.bold = True
                if item.get("duration"):
                    p.add_run(f"  ({item['duration']})")
                for bullet in (item.get("highlights") or [])[:4]:
                    bp = document.add_paragraph(style="List Bullet")
                    bp.paragraph_format.space_after = Pt(1)
                    bp.add_run(str(bullet))

        if content.projects:
            add_heading("Projects")
            for item in content.projects[:3]:
                p = document.add_paragraph()
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(1)
                r_proj = p.add_run(item.get("name", "Project"))
                r_proj.bold = True
                for bullet in (item.get("highlights") or [])[:4]:
                    bp = document.add_paragraph(style="List Bullet")
                    bp.paragraph_format.space_after = Pt(1)
                    bp.add_run(str(bullet))

        if content.education:
            add_heading("Education")
            for item in content.education[:3]:
                p = document.add_paragraph()
                p.add_run(f"{item.get('course_name', '')} — {item.get('college_name', '')}").bold = True
                if item.get("location"):
                    p.add_run(f" ({item.get('location')})")

        if content.certificates:
            add_heading("Certificates & Achievements")
            for item in content.certificates[:4]:
                bp = document.add_paragraph(style="List Bullet")
                bp.add_run(f"{item.get('certificate_name', '')} — {item.get('certificate_issuer', '')}")

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp_file:
            document.save(tmp_file.name)
            tmp_file.flush()
            tmp_file.seek(0)
            return tmp_file.read()

    # Template 2: Classic Executive ATS (Times New Roman, Centered, Upper Case Dividers)
    def _render_classic_docx(self, content: ResumeContent, user_details: dict[str, Any] | None = None) -> bytes:
        document = Document()
        for section in document.sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)

        style = document.styles["Normal"]
        style.font.name = "Times New Roman"
        style.font.size = Pt(11)

        title = document.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run(((user_details or {}).get("name") or "Resume").upper())
        run.bold = True
        run.font.size = Pt(16)

        contact = document.add_paragraph()
        contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_lines = []
        for key in ["location", "phone", "email", "linkedin", "github"]:
            val = (user_details or {}).get(key)
            if val:
                contact_lines.append(str(val))
        if contact_lines:
            c_run = contact.add_run(" • ".join(contact_lines))
            c_run.font.size = Pt(10)

        def add_heading(text):
            h = document.add_paragraph()
            h.paragraph_format.space_before = Pt(10)
            h.paragraph_format.space_after = Pt(3)
            r = h.add_run(text.upper())
            r.bold = True
            r.font.size = Pt(11.5)

        if content.summary:
            add_heading("Summary")
            document.add_paragraph(content.summary)

        if content.experience:
            add_heading("Professional Experience")
            for item in content.experience[:4]:
                p = document.add_paragraph()
                p.paragraph_format.space_before = Pt(4)
                r_role = p.add_run(f"{item.get('role', '')}")
                r_role.bold = True
                p.add_run(f", {item.get('company', '')}")
                if item.get("duration"):
                    p.add_run(f" | {item['duration']}")
                for bullet in (item.get("highlights") or [])[:4]:
                    bp = document.add_paragraph(style="List Bullet")
                    bp.add_run(str(bullet))

        if content.projects:
            add_heading("Key Projects")
            for item in content.projects[:3]:
                p = document.add_paragraph()
                p.paragraph_format.space_before = Pt(4)
                r_p = p.add_run(item.get("name", "Project"))
                r_p.bold = True
                for bullet in (item.get("highlights") or [])[:4]:
                    bp = document.add_paragraph(style="List Bullet")
                    bp.add_run(str(bullet))

        if content.skills:
            add_heading("Skills & Expertise")
            sp = document.add_paragraph()
            sp.add_run(" • ".join(content.skills[:15]))

        if content.education:
            add_heading("Education")
            for item in content.education[:3]:
                p = document.add_paragraph()
                p.add_run(f"{item.get('course_name', '')} — {item.get('college_name', '')}").bold = True

        if content.certificates:
            add_heading("Certifications")
            for item in content.certificates[:4]:
                bp = document.add_paragraph(style="List Bullet")
                bp.add_run(f"{item.get('certificate_name', '')} ({item.get('certificate_issuer', '')})")

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp_file:
            document.save(tmp_file.name)
            tmp_file.flush()
            tmp_file.seek(0)
            return tmp_file.read()

    # Template 3: Tech / Compact ATS (Helvetica / Segoe UI, Skills upfront, High density)
    def _render_tech_docx(self, content: ResumeContent, user_details: dict[str, Any] | None = None) -> bytes:
        document = Document()
        for section in document.sections:
            section.top_margin = Inches(0.35)
            section.bottom_margin = Inches(0.35)
            section.left_margin = Inches(0.4)
            section.right_margin = Inches(0.4)

        style = document.styles["Normal"]
        style.font.name = "Arial"
        style.font.size = Pt(10)

        title = document.add_paragraph()
        run = title.add_run((user_details or {}).get("name") or "Developer Resume")
        run.bold = True
        run.font.size = Pt(15)
        run.font.color.rgb = RGBColor(30, 41, 59)

        contact = document.add_paragraph()
        contact_lines = []
        for key in ["email", "phone", "github", "linkedin", "portfolio", "location"]:
            val = (user_details or {}).get(key)
            if val:
                contact_lines.append(str(val))
        if contact_lines:
            c_run = contact.add_run(" | ".join(contact_lines))
            c_run.font.size = Pt(9)
            c_run.font.color.rgb = RGBColor(100, 116, 139)

        def add_heading(text):
            h = document.add_paragraph()
            h.paragraph_format.space_before = Pt(6)
            h.paragraph_format.space_after = Pt(2)
            r = h.add_run(f"// {text.upper()}")
            r.bold = True
            r.font.size = Pt(10.5)
            r.font.color.rgb = RGBColor(37, 99, 235)

        if content.skills:
            add_heading("Core Technical Skills")
            sp = document.add_paragraph()
            sp.add_run(" | ".join(content.skills[:16])).bold = True

        if content.summary:
            add_heading("Profile Summary")
            document.add_paragraph(content.summary)

        if content.projects:
            add_heading("Featured Projects")
            for item in content.projects[:3]:
                p = document.add_paragraph()
                p.paragraph_format.space_before = Pt(3)
                p.add_run(f"{item.get('name', 'Project')}").bold = True
                for bullet in (item.get("highlights") or [])[:4]:
                    bp = document.add_paragraph(style="List Bullet")
                    bp.paragraph_format.space_after = Pt(1)
                    bp.add_run(str(bullet))

        if content.experience:
            add_heading("Experience")
            for item in content.experience[:3]:
                p = document.add_paragraph()
                p.paragraph_format.space_before = Pt(3)
                p.add_run(f"{item.get('role', '')} @ {item.get('company', '')}").bold = True
                if item.get("duration"):
                    p.add_run(f" [{item['duration']}]")
                for bullet in (item.get("highlights") or [])[:4]:
                    bp = document.add_paragraph(style="List Bullet")
                    bp.paragraph_format.space_after = Pt(1)
                    bp.add_run(str(bullet))

        if content.education:
            add_heading("Education")
            for item in content.education[:3]:
                p = document.add_paragraph()
                p.add_run(f"{item.get('course_name', '')} - {item.get('college_name', '')}").bold = True

        if content.certificates:
            add_heading("Certifications")
            for item in content.certificates[:4]:
                bp = document.add_paragraph(style="List Bullet")
                bp.add_run(f"{item.get('certificate_name', '')} ({item.get('certificate_issuer', '')})")

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp_file:
            document.save(tmp_file.name)
            tmp_file.flush()
            tmp_file.seek(0)
            return tmp_file.read()
